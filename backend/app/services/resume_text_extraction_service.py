"""Resume text extraction with a normal-text-first OCR fallback.

PDFs are parsed with pypdf first. OCR is only invoked for pages whose native
text extraction is empty or clearly insufficient. This keeps normal resumes
fast while supporting scanned/image-based PDFs and mixed PDFs.
"""
from __future__ import annotations

import io
import logging

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.config import get_settings

logger = logging.getLogger(__name__)


class ResumeTextExtractionError(ValueError):
    """Raised when a resume cannot be converted into usable text."""



def extract_resume_text(file_type: str, file_bytes: bytes) -> str:
    if file_type == "pdf":
        return extract_pdf_text(file_bytes)
    if file_type == "docx":
        return _extract_docx_text(file_bytes)
    raise ResumeTextExtractionError(f"Unsupported resume file type: {file_type}.")


def extract_pdf_text(file_bytes: bytes) -> str:
    settings = get_settings()
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
    except PdfReadError as exc:
        raise ResumeTextExtractionError("This PDF could not be read — it may be corrupted.") from exc
    except Exception as exc:
        raise ResumeTextExtractionError("This PDF could not be opened.") from exc

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            pass
        if reader.is_encrypted:
            raise ResumeTextExtractionError("This PDF is password-protected and can't be read.")

    page_texts: list[str] = []
    pages_needing_ocr: list[int] = []

    for index, page in enumerate(reader.pages):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:
            logger.warning("Native PDF text extraction failed for page %s: %s", index + 1, exc)
            text = ""

        page_texts.append(text)
        if len(_meaningful_text(text)) < settings.OCR_MIN_PAGE_TEXT_CHARS:
            pages_needing_ocr.append(index)

    if not pages_needing_ocr:
        return _join_page_text(page_texts)

    # OCR is deliberately lazy: imports and rendering only happen when at
    # least one page needs it. Missing OCR binaries/dependencies therefore do
    # not affect ordinary text PDFs.
    try:
        ocr_pages = _ocr_pdf_pages(file_bytes, pages_needing_ocr)
    except ResumeTextExtractionError:
        # If native extraction produced enough useful text overall, preserve it
        # rather than failing the upload just because OCR is unavailable.
        native_text = _join_page_text(page_texts)
        if len(_meaningful_text(native_text)) >= settings.OCR_MIN_DOCUMENT_TEXT_CHARS:
            logger.exception("OCR fallback failed; retaining native PDF text")
            return native_text
        raise

    for index, ocr_text in ocr_pages.items():
        if len(_meaningful_text(ocr_text)) > len(_meaningful_text(page_texts[index])):
            page_texts[index] = ocr_text.strip()

    combined = _join_page_text(page_texts)
    if len(_meaningful_text(combined)) < settings.OCR_MIN_DOCUMENT_TEXT_CHARS:
        raise ResumeTextExtractionError(
            "No usable text could be extracted from this PDF, including OCR. "
            "Please upload a clearer resume PDF."
        )
    return combined


def _ocr_pdf_pages(file_bytes: bytes, page_indexes: list[int]) -> dict[int, str]:
    settings = get_settings()
    try:
        try:
            import pymupdf as fitz
        except ImportError:
            import fitz  # PyMuPDF fallback
        import pytesseract
        from PIL import Image
    except ImportError as exc:
        raise ResumeTextExtractionError(
            "OCR support is not installed on the server. Install the OCR dependencies and try again."
        ) from exc

    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as document:
            results: dict[int, str] = {}
            for index in page_indexes:
                page = document.load_page(index)
                pixmap = page.get_pixmap(
                    matrix=fitz.Matrix(settings.OCR_RENDER_SCALE, settings.OCR_RENDER_SCALE),
                    alpha=False,
                )
                image = Image.open(io.BytesIO(pixmap.tobytes("png")))
                text = pytesseract.image_to_string(
                    image,
                    lang=settings.OCR_TESSERACT_LANG,
                    config=settings.OCR_TESSERACT_CONFIG,
                    timeout=settings.OCR_PAGE_TIMEOUT_SECONDS,
                )
                results[index] = text or ""
            return results
    except pytesseract.TesseractNotFoundError as exc:
        raise ResumeTextExtractionError(
            "OCR is unavailable because the Tesseract executable is not installed on the server."
        ) from exc
    except RuntimeError as exc:
        # pytesseract uses RuntimeError for OCR timeout in some versions.
        raise ResumeTextExtractionError(
            "OCR timed out while processing the resume. Please upload a shorter or clearer PDF."
        ) from exc
    except Exception as exc:
        raise ResumeTextExtractionError("OCR failed while processing the PDF.") from exc


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        document = DocxDocument(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ResumeTextExtractionError("This DOCX file could not be read — it may be corrupted.") from exc

    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    parts.append(cell.text)
    text = _join_page_text(parts)
    if not _meaningful_text(text):
        raise ResumeTextExtractionError("No usable text was found in this DOCX file.")
    return text


def _meaningful_text(text: str) -> str:
    return " ".join((text or "").split())


def _join_page_text(parts: list[str]) -> str:
    return "\n\n".join(part.strip() for part in parts if part and part.strip())
